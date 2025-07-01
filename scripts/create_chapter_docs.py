#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
C言語学習教材 章別Wordファイル生成スクリプト
"""

import os
import re
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

# 章情報の定義
CHAPTERS = [
    {"num": 1, "dir": "introduction", "title": "導入・環境構築"},
    {"num": 2, "dir": "basics-syntax", "title": "基本文法・Hello World"},
    {"num": 3, "dir": "data-types", "title": "データ型と変数"},
    {"num": 4, "dir": "operators", "title": "演算子"},
    {"num": 5, "dir": "control-if", "title": "制御構造（条件分岐）"},
    {"num": 6, "dir": "control-loop", "title": "制御構造（ループ）"},
    {"num": 7, "dir": "arrays", "title": "配列"},
    {"num": 8, "dir": "strings", "title": "文字列処理"},
    {"num": 9, "dir": "functions", "title": "関数"},
    {"num": 10, "dir": "pointers", "title": "ポインタ基礎"},
    {"num": 11, "dir": "structures", "title": "構造体とポインタ"},
    {"num": 12, "dir": "function-pointers", "title": "関数ポインタ"},
    {"num": 13, "dir": "advanced", "title": "複数ファイル・発展技術"},
    {"num": 14, "dir": "c23-features", "title": "C23の新機能（オプション）"}
]

def setup_document_styles(doc):
    """ドキュメントのスタイルを設定"""
    # タイトルスタイル
    title_style = doc.styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'メイリオ'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(18)
    
    # 見出し1スタイル
    h1_style = doc.styles.add_style('Heading1Custom', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = 'メイリオ'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 139)  # Dark Blue
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    
    # 見出し2スタイル
    h2_style = doc.styles.add_style('Heading2Custom', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = 'メイリオ'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.paragraph_format.space_before = Pt(6)
    h2_style.paragraph_format.space_after = Pt(3)
    
    # コードスタイル
    code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(3)
    
    # 本文スタイル
    body_style = doc.styles['Normal']
    body_style.font.name = 'メイリオ'
    body_style.font.size = Pt(11)
    body_style.paragraph_format.line_spacing = 1.5

def read_markdown_file(filepath):
    """Markdownファイルを読み込む"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return ""

def read_c_file(filepath):
    """Cソースファイルを読み込む"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return ""

def convert_markdown_to_docx(doc, markdown_content, chapter_info):
    """Markdownコンテンツを処理してDocxに変換"""
    lines = markdown_content.split('\n')
    in_code_block = False
    code_buffer = []
    
    # タイトルを追加
    title = doc.add_paragraph(f"第{chapter_info['num']}章 {chapter_info['title']}", style='ChapterTitle')
    
    for line in lines:
        # コードブロックの処理
        if line.strip().startswith('```'):
            if in_code_block:
                # コードブロック終了
                if code_buffer:
                    code_text = '\n'.join(code_buffer)
                    code_para = doc.add_paragraph(code_text, style='CodeBlock')
                code_buffer = []
                in_code_block = False
            else:
                # コードブロック開始
                in_code_block = True
            continue
        
        if in_code_block:
            code_buffer.append(line)
            continue
        
        # 見出しの処理
        if line.startswith('## '):
            doc.add_paragraph(line[3:], style='Heading1Custom')
        elif line.startswith('### '):
            doc.add_paragraph(line[4:], style='Heading2Custom')
        elif line.startswith('# '):
            # 章タイトルはスキップ（既に追加済み）
            continue
        elif line.strip():
            # 通常のテキスト
            # インラインコードを処理
            line = re.sub(r'`([^`]+)`', r'\1', line)
            # 太字を処理
            line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
            # リンクを処理
            line = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line)
            
            # リスト項目の処理
            if line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(line, style='List Number')
            else:
                doc.add_paragraph(line)

def create_textbook_doc(chapter_info, src_dir, output_dir):
    """教科書用Docxファイルを作成"""
    doc = Document()
    setup_document_styles(doc)
    
    # READMEの内容を追加
    readme_path = os.path.join(src_dir, chapter_info['dir'], 'README.md')
    readme_content = read_markdown_file(readme_path)
    
    if readme_content:
        convert_markdown_to_docx(doc, readme_content, chapter_info)
    
    # サンプルコードを追加
    examples_dir = os.path.join(src_dir, chapter_info['dir'], 'examples')
    if os.path.exists(examples_dir):
        doc.add_page_break()
        doc.add_paragraph('サンプルコード', style='Heading1Custom')
        
        for filename in sorted(os.listdir(examples_dir)):
            if filename.endswith('.c') and not filename.endswith('_c99.c'):
                filepath = os.path.join(examples_dir, filename)
                code_content = read_c_file(filepath)
                
                if code_content:
                    doc.add_paragraph(f"\n{filename}", style='Heading2Custom')
                    doc.add_paragraph(code_content, style='CodeBlock')
    
    # ファイルを保存
    output_path = os.path.join(output_dir, f"chapter_{chapter_info['num']:02d}_textbook.docx")
    doc.save(output_path)
    print(f"作成: {output_path}")

def create_exercises_doc(chapter_info, src_dir, output_dir):
    """演習問題用Docxファイルを作成"""
    doc = Document()
    setup_document_styles(doc)
    
    # タイトルを追加
    title = doc.add_paragraph(f"第{chapter_info['num']}章 {chapter_info['title']} - 演習問題", style='ChapterTitle')
    
    # 演習問題を追加
    exercises_readme = os.path.join(src_dir, chapter_info['dir'], 'exercises', 'README.md')
    exercises_content = read_markdown_file(exercises_readme)
    
    if exercises_content:
        doc.add_paragraph('演習問題', style='Heading1Custom')
        # 演習問題のREADMEから内容を抽出
        lines = exercises_content.split('\n')
        for line in lines:
            if line.strip() and not line.startswith('#'):
                if line.startswith('##'):
                    doc.add_paragraph(line[2:].strip(), style='Heading2Custom')
                else:
                    doc.add_paragraph(line)
    
    # 解答例を追加
    solutions_dir = os.path.join(src_dir, chapter_info['dir'], 'solutions')
    if os.path.exists(solutions_dir):
        doc.add_page_break()
        doc.add_paragraph('解答例', style='Heading1Custom')
        
        # 解答例のREADMEを追加
        solutions_readme = os.path.join(solutions_dir, 'README.md')
        solutions_content = read_markdown_file(solutions_readme)
        
        if solutions_content:
            lines = solutions_content.split('\n')
            for line in lines:
                if line.strip() and not line.startswith('#'):
                    doc.add_paragraph(line)
        
        # 解答コードを追加
        for filename in sorted(os.listdir(solutions_dir)):
            if filename.endswith('.c') and not filename.endswith('_c99.c'):
                filepath = os.path.join(solutions_dir, filename)
                code_content = read_c_file(filepath)
                
                if code_content:
                    doc.add_paragraph(f"\n{filename}", style='Heading2Custom')
                    doc.add_paragraph(code_content, style='CodeBlock')
    
    # ファイルを保存
    output_path = os.path.join(output_dir, f"chapter_{chapter_info['num']:02d}_exercises.docx")
    doc.save(output_path)
    print(f"作成: {output_path}")

def main():
    """メイン処理"""
    src_dir = '../src'
    output_dir = '../output/word'
    
    # 出力ディレクトリの作成
    os.makedirs(output_dir, exist_ok=True)
    
    print(f"C言語学習教材 章別Wordファイル生成")
    print(f"生成日時: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"出力先: {output_dir}")
    print("-" * 50)
    
    # 各章のファイルを生成
    for chapter in CHAPTERS:
        print(f"\n第{chapter['num']}章: {chapter['title']}")
        
        # 教科書ファイルを作成
        try:
            create_textbook_doc(chapter, src_dir, output_dir)
        except Exception as e:
            print(f"  エラー (教科書): {e}")
        
        # 演習問題ファイルを作成
        try:
            create_exercises_doc(chapter, src_dir, output_dir)
        except Exception as e:
            print(f"  エラー (演習問題): {e}")
    
    print("\n" + "-" * 50)
    print("完了しました。")

if __name__ == "__main__":
    main()