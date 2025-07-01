#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
章別DOCXファイルを統合して教科書と演習問題のDOCXを作成
"""

import os
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_break(doc):
    """ページブレークを追加"""
    doc.add_page_break()

def copy_element(element, target_doc):
    """要素をターゲットドキュメントにコピー"""
    # 段落をコピー
    if element.style.name.startswith('Heading') or element.style.name == 'Title':
        new_para = target_doc.add_paragraph(element.text, style=element.style.name)
    else:
        new_para = target_doc.add_paragraph()
        new_para.alignment = element.alignment
        
        # ランをコピー
        for run in element.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb

def merge_docx_files(file_list, output_path, title):
    """複数のDOCXファイルを1つに統合"""
    print(f"\n{title}を作成中...")
    
    # 最初のファイルを基準として使用
    if not file_list:
        print("統合するファイルがありません")
        return
    
    # 新しいドキュメントを作成（最初のファイルのスタイルを継承）
    base_doc = Document(file_list[0])
    merged_doc = Document()
    
    # カスタムスタイルをコピー
    for style in base_doc.styles:
        try:
            if style.name not in [s.name for s in merged_doc.styles]:
                merged_doc.styles.add_style(style.name, style.type)
        except:
            pass
    
    # タイトルページを追加
    title_para = merged_doc.add_paragraph(title, style='Title')
    add_page_break(merged_doc)
    
    # 各ファイルを処理
    for idx, docx_file in enumerate(file_list):
        print(f"  処理中: {os.path.basename(docx_file)}")
        
        try:
            doc = Document(docx_file)
            
            # 各段落をコピー
            for element in doc.paragraphs:
                if element.text.strip():  # 空の段落をスキップ
                    try:
                        copy_element(element, merged_doc)
                    except:
                        # スタイルが存在しない場合は通常の段落として追加
                        merged_doc.add_paragraph(element.text)
            
            # 表をコピー
            for table in doc.tables:
                # 新しい表を作成
                new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                new_table.style = table.style
                
                # セルの内容をコピー
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.rows[i].cells[j].text = cell.text
            
            # 最後のファイル以外は改ページを追加
            if idx < len(file_list) - 1:
                add_page_break(merged_doc)
                
        except Exception as e:
            print(f"  エラー: {docx_file} - {str(e)}")
    
    # ファイルを保存
    merged_doc.save(output_path)
    print(f"  完成: {output_path}")

def main():
    """メイン処理"""
    docs_dir = "../output/word"
    
    # 教科書ファイルのリスト
    textbook_files = []
    exercises_files = []
    
    # ファイルを収集
    for i in range(1, 15):  # 1-14章
        textbook_file = os.path.join(docs_dir, f"chapter_{i:02d}_textbook.docx")
        exercises_file = os.path.join(docs_dir, f"chapter_{i:02d}_exercises.docx")
        
        if os.path.exists(textbook_file):
            textbook_files.append(textbook_file)
        else:
            print(f"警告: {textbook_file} が見つかりません")
            
        if os.path.exists(exercises_file):
            exercises_files.append(exercises_file)
        else:
            print(f"警告: {exercises_file} が見つかりません")
    
    print(f"統合対象:")
    print(f"  教科書ファイル: {len(textbook_files)}個")
    print(f"  演習問題ファイル: {len(exercises_files)}個")
    
    # 教科書を統合
    if textbook_files:
        output_textbook = os.path.join(docs_dir, "教科書.DOCX")
        merge_docx_files(textbook_files, output_textbook, "C言語プログラミング教科書")
    
    # 演習問題を統合
    if exercises_files:
        output_exercises = os.path.join(docs_dir, "演習問題.DOCX")
        merge_docx_files(exercises_files, output_exercises, "C言語プログラミング演習問題集")
    
    print("\n統合完了")

if __name__ == "__main__":
    main()